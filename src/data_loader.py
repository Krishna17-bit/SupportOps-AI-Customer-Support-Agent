from __future__ import annotations

import json
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Iterable, List, Optional

import pandas as pd


@dataclass
class KnowledgeDoc:
    title: str
    content: str
    source_type: str = "document"
    metadata: Optional[dict] = None


def _read_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        pages = []
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {idx + 1}]\n{text}")
        return "\n\n".join(pages)
    except Exception as exc:  # pragma: no cover - depends on optional file quirks
        return f"Could not parse PDF safely: {exc}"


def _read_docx(raw: bytes) -> str:
    try:
        from docx import Document

        doc = Document(BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    tables.append(" | ".join(cells))
        return "\n".join(paragraphs + tables)
    except Exception as exc:  # pragma: no cover
        return f"Could not parse DOCX safely: {exc}"


def _read_tabular(raw: bytes, suffix: str) -> str:
    try:
        if suffix == ".csv":
            df = pd.read_csv(BytesIO(raw))
        else:
            df = pd.read_excel(BytesIO(raw))
        return df.head(250).to_markdown(index=False)
    except Exception as exc:
        return f"Could not parse table safely: {exc}"


def load_knowledge_uploads(uploaded_files: Optional[Iterable]) -> List[KnowledgeDoc]:
    docs: List[KnowledgeDoc] = []
    if not uploaded_files:
        return docs

    for file in uploaded_files:
        name = getattr(file, "name", "uploaded_file")
        suffix = Path(name).suffix.lower()
        raw = file.read()

        if suffix == ".pdf":
            content = _read_pdf(raw)
        elif suffix == ".docx":
            content = _read_docx(raw)
        elif suffix in {".csv", ".xlsx", ".xls"}:
            content = _read_tabular(raw, suffix)
        elif suffix == ".json":
            try:
                content = json.dumps(json.loads(raw.decode("utf-8")), indent=2)
            except Exception:
                content = raw.decode("utf-8", errors="ignore")
        else:
            content = raw.decode("utf-8", errors="ignore")

        docs.append(
            KnowledgeDoc(
                title=name,
                content=content[:120_000],
                source_type=suffix.replace(".", "") or "text",
                metadata={"filename": name},
            )
        )
    return docs


def load_sample_knowledge(base_dir: Path) -> List[KnowledgeDoc]:
    sample_dir = base_dir / "sample_data"
    docs: List[KnowledgeDoc] = []
    for path in sorted(sample_dir.glob("*")):
        if path.name == "tickets.csv":
            continue
        if path.suffix.lower() in {".md", ".txt"}:
            docs.append(KnowledgeDoc(path.name, path.read_text(encoding="utf-8"), path.suffix[1:]))
        elif path.suffix.lower() == ".csv":
            try:
                df = pd.read_csv(path)
                docs.append(KnowledgeDoc(path.name, df.to_markdown(index=False), "csv"))
            except Exception:
                docs.append(KnowledgeDoc(path.name, path.read_text(encoding="utf-8", errors="ignore"), "csv"))
    return docs


def load_tickets(uploaded_ticket_file, base_dir: Path) -> pd.DataFrame:
    if uploaded_ticket_file is not None:
        raw = uploaded_ticket_file.read()
        name = getattr(uploaded_ticket_file, "name", "tickets.csv")
        suffix = Path(name).suffix.lower()
        if suffix == ".csv":
            df = pd.read_csv(BytesIO(raw))
        elif suffix in {".xlsx", ".xls"}:
            df = pd.read_excel(BytesIO(raw))
        elif suffix == ".json":
            df = pd.read_json(BytesIO(raw))
        else:
            text = raw.decode("utf-8", errors="ignore")
            rows = []
            for idx, block in enumerate(text.split("\n\n"), start=1):
                if block.strip():
                    rows.append({"ticket_id": f"TXT-{idx:03d}", "subject": block[:80], "body": block})
            df = pd.DataFrame(rows)
    else:
        df = pd.read_csv(base_dir / "sample_data" / "tickets.csv")

    return normalize_ticket_columns(df)


def normalize_ticket_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    clean_cols = {c: str(c).strip().lower().replace(" ", "_") for c in df.columns}
    df.rename(columns=clean_cols, inplace=True)

    aliases = {
        "id": "ticket_id",
        "ticket": "ticket_id",
        "customer": "customer_name",
        "name": "customer_name",
        "message": "body",
        "description": "body",
        "issue": "body",
        "title": "subject",
        "email_address": "email",
        "order": "order_id",
        "created": "created_at",
    }
    for old, new in aliases.items():
        if old in df.columns and new not in df.columns:
            df[new] = df[old]

    required = ["ticket_id", "customer_name", "email", "subject", "body", "channel", "created_at", "order_id"]
    for col in required:
        if col not in df.columns:
            if col == "ticket_id":
                df[col] = [f"T-{i + 1:04d}" for i in range(len(df))]
            elif col == "created_at":
                df[col] = pd.Timestamp.today().strftime("%Y-%m-%d %H:%M")
            else:
                df[col] = ""

    df["subject"] = df["subject"].fillna("").astype(str)
    df["body"] = df["body"].fillna("").astype(str)
    df["full_text"] = (df["subject"] + "\n" + df["body"]).str.strip()
    return df


def load_order_history(base_dir: Path) -> pd.DataFrame:
    path = base_dir / "sample_data" / "order_history.csv"
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame(columns=["order_id", "order_date", "customer_email", "amount", "status", "delivered_date"])
